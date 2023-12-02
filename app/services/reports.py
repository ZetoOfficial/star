from datetime import datetime
from uuid import UUID

from docx import Document
from openpyxl import Workbook
from openpyxl.styles import Font, Border, Side
from openpyxl.utils import get_column_letter

from app.repositories import ReportRepository
from app.schemas import ReportDTO, GalaxyReportDTO
from settings import load_settings

sttngs = load_settings()


# Constants
COLUMN_TITLES = ["Категория", "Описание"]
COLUMN_WIDTH_PADDING = 2


def get_public_url(file_path: str) -> str:
    """Generate a public URL for a given file path."""
    return f"/{file_path.lstrip('/')}"


def format_report_data(report_data):
    """Format report data for both Excel and Word reports."""
    galaxy = report_data["galaxy"]
    stars = [star.name for star in report_data["stars"]]
    planets = [planet.name for planet in report_data["planets"]]

    return [
        ["Название галактики", galaxy.name],
        ["Тип галактики", galaxy.shape],
        ["Размер галактики", f"{galaxy.size} млн. км"],
        ["Основные звезды", ", ".join(stars) if stars else "Звезды не найдены"],
        ["Известные планеты", ", ".join(planets) if planets else "Планеты не найдены"],
        ["Состав", galaxy.composition],
    ]


class ReportsService:
    @staticmethod
    async def get_report_data(galaxy_id: UUID):
        report_data = await ReportRepository.get_report_data(galaxy_id)
        stars = [star.name for star in report_data["stars"]]
        planets = [planet.name for planet in report_data["planets"]]
        return GalaxyReportDTO(
            name=report_data["galaxy"].name,
            shape=report_data["galaxy"].shape,
            size=report_data["galaxy"].size,
            stars=", ".join(stars) if stars else "Звезды не найдены",
            planets=", ".join(planets) if planets else "Планеты не найдены",
            composition=report_data["galaxy"].composition,
        )

    @staticmethod
    async def get_excel_data(galaxy_id: UUID):
        report_data = await ReportRepository.get_report_data(galaxy_id)
        formatted_data = format_report_data(report_data)

        wb = Workbook()
        ws = wb.active
        ws.title = "Отчет о Галактике"

        # Styles
        bold_font = Font(bold=True)
        thin_border = Border(
            left=Side(style="thin"),
            right=Side(style="thin"),
            top=Side(style="thin"),
            bottom=Side(style="thin"),
        )

        # Add headers
        ws.append(COLUMN_TITLES)
        for cell in ws[1]:
            cell.font = bold_font
            cell.border = thin_border

        # Add data rows
        for row in formatted_data:
            ws.append(row)
            for cell in ws[ws.max_row]:
                cell.border = thin_border

        # Auto-adjust column widths
        for column in ws.columns:
            max_length = max(len(str(cell.value)) for cell in column if cell.value)
            adjusted_width = max_length + COLUMN_WIDTH_PADDING
            ws.column_dimensions[
                get_column_letter(column[0].column)
            ].width = adjusted_width

        # Save the workbook
        report_name = f"{sttngs.app.static_dir}/report_{report_data['galaxy'].name}_{datetime.now().strftime('%Y-%m-%d_%H-%M-%S')}.xlsx"
        wb.save(report_name)

        return ReportDTO(name=report_name, link=get_public_url(report_name))

    @staticmethod
    async def get_word_data(galaxy_id: UUID):
        report_data = await ReportRepository.get_report_data(galaxy_id)
        formatted_data = format_report_data(report_data)

        doc = Document()
        doc.add_heading("Отчет о Галактике", level=1)

        # Add a table
        table = doc.add_table(rows=1, cols=2)
        table.style = "Table Grid"

        # Add header row
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text, hdr_cells[1].text = COLUMN_TITLES
        for cell in hdr_cells:
            cell.paragraphs[0].runs[0].font.bold = True

        # Add data rows
        for category, description in formatted_data:
            row_cells = table.add_row().cells
            row_cells[0].text, row_cells[1].text = category, description

        # Save the document
        report_name = f"{sttngs.app.static_dir}/report_{report_data['galaxy'].name}_{datetime.now().strftime('%Y-%m-%d_%H-%M-%S')}.docx"
        doc.save(report_name)

        return ReportDTO(name=report_name, link=get_public_url(report_name))
